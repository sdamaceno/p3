import streamlit as st # type: ignore[import]
import pandas as pd
import concurrent.futures
import time
import io
import re
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta, timezone
from dateutil.relativedelta import relativedelta

from utils import formatar_moeda_simples, formatar_moeda_ordenavel, gerar_hash_item, parse_numero_localizado
from estatistica import processar_precos_regra, ordenar_validos, ordenar_outliers
from pncp_api import PNCPEngine
from pdf_engine import gerar_pdf_oficial
import typography
import layout
import components
from texts import get_text

st.set_page_config(page_title=get_text("NOME_SISTEMA"), layout="wide", initial_sidebar_state="collapsed")
fuso_br = timezone(timedelta(hours=-3))

typography.load()
layout.load()
components.load()

# --- MECANISMO DE SERIALIZAÇÃO XML (IMPORTAR/EXPORTAR) ---
def exportar_contratacao_xml(ctx):
    root = ET.Element("contratacao")
    ET.SubElement(root, "titulo").text = str(ctx.get("titulo", ""))
    ET.SubElement(root, "objeto_contratacao").text = str(ctx.get("objeto_contratacao", ""))
    ET.SubElement(root, "regra_calculo").text = str(ctx.get("regra_calculo", ""))
    ET.SubElement(root, "meses_corte").text = str(ctx.get("meses_corte", "24"))
    ET.SubElement(root, "paginas_pncp").text = str(ctx.get("paginas_pncp", "3"))
    
    delay = ctx.get("delay_pncp")
    ET.SubElement(root, "delay_pncp").text = str(delay) if delay is not None else ""
    
    el_itens = ET.SubElement(root, "itens_tr")
    df_tr = ctx.get("df_tr", pd.DataFrame())
    for _, row in df_tr.iterrows():
        el_item = ET.SubElement(el_itens, "item_tr")
        ET.SubElement(el_item, "Lote").text = str(row.get("Lote", ""))
        ET.SubElement(el_item, "Item").text = str(row.get("Item", ""))
        ET.SubElement(el_item, "Descricao").text = str(row.get("Descrição", ""))
        ET.SubElement(el_item, "Metrica").text = str(row.get("Métrica", ""))
        ET.SubElement(el_item, "Tipo").text = str(row.get("Tipo", ""))
        ET.SubElement(el_item, "Quantidade").text = str(row.get("Quantidade", ""))
        
        h = gerar_hash_item(row)
        kw = ctx.get("palavras_chave_massa", {}).get(h, "")
        ET.SubElement(el_item, "PalavraChave").text = str(kw)
        
        banco = ctx.get("banco_precos", {}).get(h, {})
        if banco:
            el_banco = ET.SubElement(el_item, "banco_precos")
            ET.SubElement(el_banco, "estatistica_pronta").text = str(banco.get("estatistica_pronta", False))
            ET.SubElement(el_banco, "media_saneada").text = str(banco.get("media_saneada", 0.0))
            ET.SubElement(el_banco, "mediana").text = str(banco.get("mediana", 0.0))
            ET.SubElement(el_banco, "amostras").text = str(banco.get("amostras", 0))
            
            for df_key in ["df_pncp", "df_manual_rastreio", "historico_buscas"]:
                df_data = banco.get(df_key, pd.DataFrame())
                el_df = ET.SubElement(el_banco, df_key)
                if not df_data.empty:
                    for _, r in df_data.iterrows():
                        el_row = ET.SubElement(el_df, "row")
                        for col in df_data.columns:
                            el_col = ET.SubElement(el_row, "field")
                            el_col.set("name", str(col))
                            el_col.text = str(r[col])
                            
    return ET.tostring(root, encoding="utf-8")

def importar_contratacao_xml(xml_file, current_id):
    try:
        xml_data = xml_file.getvalue()
        root = ET.fromstring(xml_data)
        
        ctx = {
            "id": current_id,
            "titulo": root.find("titulo").text or f"Contratação {current_id}",
            "objeto_contratacao": root.find("objeto_contratacao").text or "",
            "regra_calculo": root.find("regra_calculo").text or "Preços válidos - Mediana ±25% e Média",
            "meses_corte": int(root.find("meses_corte").text or "24"),
            "paginas_pncp": int(root.find("paginas_pncp").text or "3"),
            "tr_objeto_salvo": False,
            "tr_itens_salvos": False,
            "banco_precos": {},
            "palavras_chave_massa": {},
            "acao_ativa": (None, None),
            "massa_status": "idle",
            "massa_idx": 0,
            "massa_total": 0,
            "massa_novos": 0
        }
        
        delay_text = root.find("delay_pncp").text
        ctx["delay_pncp"] = int(delay_text) if delay_text else None
        
        if ctx["objeto_contratacao"]:
            ctx["tr_objeto_salvo"] = True
            
        itens_list = []
        el_itens = root.find("itens_tr")
        if el_itens is not None:
            for el_item in el_itens.findall("item_tr"):
                item_data = {
                    "Lote": el_item.find("Lote").text or "",
                    "Item": el_item.find("Item").text or "",
                    "Descrição": el_item.find("Descricao").text or "",
                    "Métrica": el_item.find("Metrica").text or "",
                    "Tipo": el_item.find("Tipo").text or "",
                    "Quantidade": parse_numero_localizado(el_item.find("Quantidade").text or "0")
                }
                itens_list.append(item_data)
                
                row_sim = pd.Series(item_data)
                h = gerar_hash_item(row_sim)
                
                kw = el_item.find("PalavraChave").text or ""
                ctx["palavras_chave_massa"][h] = kw
                
                el_banco = el_item.find("banco_precos")
                if el_banco is not None:
                    banco_item = {
                        "estatistica_pronta": el_banco.find("estatistica_pronta").text == "True",
                        "media_saneada": float(el_banco.find("media_saneada").text or "0.0"),
                        "mediana": float(el_banco.find("mediana").text or "0.0"),
                        "amostras": int(el_banco.find("amostras").text or "0"),
                        "df_validos": pd.DataFrame(),
                        "df_outliers": pd.DataFrame()
                    }
                    
                    for df_key in ["df_pncp", "df_manual_rastreio", "historico_buscas"]:
                        el_df = el_banco.find(df_key)
                        rows_data = []
                        if el_df is not None:
                            for el_row in el_df.findall("row"):
                                r_dict = {}
                                for el_col in el_row.findall("field"):
                                    col_name = el_col.get("name")
                                    val = el_col.text or ""
                                    if val == "True": val = True
                                    elif val == "False": val = False
                                    else:
                                        try:
                                            if "." in val: val = float(val)
                                            else: val = int(val)
                                        except ValueError:
                                            pass
                                    r_dict[col_name] = val
                                rows_data.append(r_dict)
                        
                        if df_key == "df_pncp":
                            cols_def = ["Válido?", "Data", "Empresa/Órgão", "Item", "Qtd", "Preço", "Valor Unitário", "Origem", "Tipo"]
                        elif df_key == "df_manual_rastreio":
                            cols_def = ["Data do Contato", "Horário", "Empresa", "CNPJ/CPF", "Tipo de fonte", "Descrição da fonte", "Link da fonte", "Nome do Contato", "E-mail", "Telefone", "Situação", "Preço", "Valor Unitário"]
                        else:
                            cols_def = ["Data/Hora", "Termo Pesquisado", "Novos Registros"]
                            
                        banco_item[df_key] = pd.DataFrame(rows_data) if rows_data else pd.DataFrame(columns=cols_def)
                        
                    ctx["banco_precos"][h] = banco_item
                else:
                    ctx["banco_precos"][h] = {"df_pncp": pd.DataFrame(columns=["Válido?", "Data", "Empresa/Órgão", "Item", "Qtd", "Preço", "Valor Unitário", "Origem", "Tipo"]), "df_manual_rastreio": pd.DataFrame(columns=["Data do Contato", "Horário", "Empresa", "CNPJ/CPF", "Tipo de fonte", "Descrição da fonte", "Link da fonte", "Nome do Contato", "E-mail", "Telefone", "Situação", "Preço", "Valor Unitário"]), "historico_buscas": pd.DataFrame(columns=["Data/Hora", "Termo Pesquisado", "Novos Registros"]), "estatistica_pronta": False, "media_saneada": 0.0, "mediana": 0.0, "amostras": 0, "df_validos": pd.DataFrame(), "df_outliers": pd.DataFrame()}
        
        if itens_list:
            ctx["df_tr"] = pd.DataFrame(itens_list)
            ctx["tr_itens_salvos"] = True
        else:
            ctx["df_tr"] = pd.DataFrame(columns=["Lote", "Item", "Descrição", "Métrica", "Tipo", "Quantidade"])
            
        return ctx
    except Exception as e:
        st.error(f"Erro estrutural ao processar o arquivo XML: {e}")
        return None

# Motor Avançado de Extração Semântica e Estrutural (Smart Extract)
def processar_arquivo_inteligente(uploaded_file):
    ext = uploaded_file.name.split('.')[-1].lower()
    texto = ""
    df_final = pd.DataFrame()
    objeto_str = ""
    try:
        if ext in ['xlsx', 'xls', 'ods']: 
            df_final = pd.read_excel(uploaded_file)
        elif ext in ['csv']: 
            df_final = pd.read_csv(uploaded_file, delimiter=';')
        elif ext == 'pdf':
            import pdfplumber
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    texto += page.extract_text() + "\n"
                    tabs = page.extract_tables()
                    if tabs and df_final.empty:
                        for t in tabs:
                            if t and len(t) > 1 and len(t[0]) >= 4:
                                cols = [str(c).strip() if c else f"Col_{idx}" for idx, c in enumerate(t[0])]
                                df_final = pd.DataFrame(t[1:], columns=cols)
                                break
        elif ext in ['docx', 'doc']:
            import docx
            doc = docx.Document(uploaded_file)
            texto = "\n".join([p.text for p in doc.paragraphs])
            if doc.tables and df_final.empty:
                t = doc.tables[0]
                data = [[cell.text.strip() for cell in row.cells] for row in t.rows]
                if len(data) > 1:
                    cols = [str(c).strip() if c else f"Col_{idx}" for idx, c in enumerate(data[0])]
                    df_final = pd.DataFrame(data[1:], columns=cols)
        elif ext == 'odt':
            import zipfile
            import xml.etree.ElementTree as ET
            with zipfile.ZipFile(uploaded_file) as z:
                content_xml = z.read('content.xml')
                root = ET.fromstring(content_xml)
                paragraphs = []
                for el in root.iter():
                    if el.tag.endswith('}p') or el.tag.endswith('}h'):
                        txt = "".join(el.itertext()).strip()
                        if txt: 
                            paragraphs.append(txt)
                texto = "\n".join(paragraphs)
                
                tabelas_odt = []
                for table in root.iter('{urn:oasis:names:tc:opendocument:xmlns:table:1.0}table'):
                    rows_data = []
                    for row in table.iter('{urn:oasis:names:tc:opendocument:xmlns:table:1.0}table-row'):
                        row_cells = []
                        for cell in row.iter('{urn:oasis:names:tc:opendocument:xmlns:table:1.0}table-cell'):
                            cell_text = "".join(cell.itertext()).strip()
                            row_cells.append(cell_text)
                        if any(row_cells): 
                            rows_data.append(row_cells)
                    if len(rows_data) > 1:
                        max_cols = max(len(r) for r in rows_data)
                        for r in rows_data:
                            while len(r) < max_cols: 
                                r.append("")
                        cols = [str(c).strip() if c else f"Col_{idx}" for idx, c in enumerate(rows_data[0])]
                        df_potencial = pd.DataFrame(rows_data[1:], columns=cols)
                        if len(df_potencial.columns) >= 4:
                            tabelas_odt.append(df_potencial)
                if tabelas_odt and df_final.empty:
                    df_final = tabelas_odt[0]
        elif ext in ['md', 'txt']:
            texto = uploaded_file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        st.error(f"Erro no processador de arquivos do Design System: {e}")

    if texto:
        match = re.search(r'(?i)1\.\s*OBJETO\s*(.*?)(?=\n\s*[2-9]\.\s|ESTRUTURA DE LOTES|Lote\s|Item\s|2\.\s|3\.\s|$)', texto, re.DOTALL)
        objeto_str = match.group(1).strip() if match else texto[:1500]
        
        if df_final.empty and '\t' in texto:
            linhas = texto.split('\n')
            for i, l in enumerate(linhas):
                if '\t' in l:
                    try:
                        tabela_str = "\n".join(linhas[i:]).strip()
                        df_raw = pd.read_csv(io.StringIO(tabela_str), sep='\t', header=None, dtype=str)
                        df_final = df_raw.fillna("").apply(lambda x: x.str.strip().str.replace('"', ''))
                    except: 
                        pass
                    break

    if not df_final.empty:
        colunas_oficiais = ["Lote", "Item", "Descrição", "Métrica", "Tipo", "Quantidade"]
        mapa_chaves = {"lote": "Lote", "item": "Item", "descri": "Descrição", "métric": "Métrica", "metric": "Métrica", "tipo": "Tipo", "quant": "Quantidade", "qtd": "Quantidade"}
        try:
            df_final.columns = df_final.columns.astype(str).str.strip()
            primeira_linha = df_final.columns.str.lower()
            
            if any(any(chave in celula for chave in mapa_chaves.keys()) for celula in primeira_linha):
                df_mapeado = pd.DataFrame()
                for k, v in mapa_chaves.items():
                    for col in df_final.columns:
                        if k in str(col).lower():
                            df_mapeado[v] = df_final[col]
                            break
                for col in colunas_oficiais:
                    if col not in df_mapeado.columns: 
                        df_mapeado[col] = ""
                df_final = df_mapeado[colunas_oficiais]
            else:
                primeira_linha_dados = df_final.iloc[0].astype(str).str.lower().str.strip()
                if any(any(chave in celula for chave in mapa_chaves.keys()) for celula in primeira_linha_dados):
                    indice_mapeado = {mapa_chaves[k]: idx for idx, cel in enumerate(primeira_linha_dados) for k in mapa_chaves if k in cel}
                    dados_uteis = df_final.iloc[1:].reset_index(drop=True)
                    df_mapeado = pd.DataFrame(columns=colunas_oficiais)
                    for col in colunas_oficiais: 
                        df_mapeado[col] = dados_uteis.iloc[:, indice_mapeado[col]] if col in indice_mapeado else ""
                    df_final = df_mapeado
                else:
                    df_pos = pd.DataFrame(columns=colunas_oficiais)
                    for i, col in enumerate(colunas_oficiais):
                        df_pos[col] = df_final.iloc[:, i] if i < len(df_final.columns) else ""
                    df_final = df_pos
        except: 
            pass

        df_final['Lote'] = df_final.get('Lote', pd.Series(dtype=str)).replace(r'^\s*$', pd.NA, regex=True).ffill()
        df_final['Quantidade'] = df_final.get('Quantidade', pd.Series(dtype=str)).apply(parse_numero_localizado)
        df_final = df_final.dropna(subset=["Item", "Descrição"]).copy()

    return objeto_str, df_final

cols_pncp = ["Válido?", "Data", "Empresa/Órgão", "Item", "Qtd", "Preço", "Valor Unitário", "Origem", "Tipo"]
cols_rastreio = ["Data do Contato", "Horário", "Empresa", "CNPJ/CPF", "Tipo de fonte", "Descrição da fonte", "Link da fonte", "Nome do Contato", "E-mail", "Telefone", "Situação", "Preço", "Valor Unitário"]
cols_historico_busca = ["Data/Hora", "Termo Pesquisado", "Novos Registros"]
opcoes_origem_decreto = ["VI - Pesquisa direta c/ fornecedores", "I - Base estadual NFe", "II - Portal de Compras GO", "III - PNCP / Ferramentas específicas", "IV - Mídia / Tabelas / Sítios eletrônicos", "V - Contratações similares da adm. pública"]
opcoes_situacao = ["Solicitação de proposta enviada", "Confirmação de recebimento da solicitação", "Proposta recebida", "Não enviou proposta comercial", "Proposta recebida com equívoco", "Proposta retificada recebida"]

CONFIG_COLUNAS_TABELA = {
    "Lote": st.column_config.TextColumn("Lote", width=55),
    "Item": st.column_config.TextColumn("Item", width=55),
    "Descrição": st.column_config.TextColumn("Descrição do Objeto", required=True, width=450),
    "Métrica": st.column_config.TextColumn("Métrica", width=85),
    "Tipo": st.column_config.TextColumn("Tipo", width=95),
    "Quantidade": st.column_config.NumberColumn("Quantidade", width=85, format="%g"),
}

if 'lista_contratacoes' not in st.session_state:
    st.session_state['lista_contratacoes'] = [{ "id": 1, "titulo": "Contratação 1", "tr_objeto_salvo": False, "tr_itens_salvos": False, "objeto_contratacao": "", "df_tr": pd.DataFrame(columns=["Lote", "Item", "Descrição", "Métrica", "Tipo", "Quantidade"]), "banco_precos": {}, "acao_ativa": (None, None), "regra_calculo": "Preços válidos - Mediana ±25% e Média", "meses_corte": 24, "paginas_pncp": 3, "delay_pncp": None, "palavras_chave_massa": {}, "massa_status": "idle", "massa_idx": 0, "massa_total": 0, "massa_novos": 0 }]

for ctx in st.session_state['lista_contratacoes']:
    if 'massa_status' not in ctx:
        ctx['massa_status'] = 'idle'
        ctx['massa_idx'] = 0
        ctx['massa_total'] = 0
        ctx['massa_novos'] = 0

st.markdown(f"<h1>{get_text('NOME_SISTEMA')}</h1>", unsafe_allow_html=True)

# CSS Unificado inline para as fontes reais dos labels customizados
ESTILO_LABEL_PADRAO = "<label style='font-family: system-ui, -apple-system, BlinkMacSystemFont, \"Segoe UI\", Roboto, sans-serif !important; font-size: 12pt !important; font-weight: 400 !important; color: #1E293B !important; margin-bottom: 6px; display: block;'>"

for idx_c, ctx in enumerate(st.session_state['lista_contratacoes']):
    with st.expander(ctx['titulo'], expanded=(idx_c == len(st.session_state['lista_contratacoes']) - 1)):
        
        # --- PAINEL INTEGRADO: RENOVAÇÃO DE TÍTULO EMBUTIDO E DATA MANAGEMENT XML ---
        c_tit, c_btn_tit, c_xml_imp, c_xml_exp = st.columns([2, 1, 1, 1], gap="small")
        with c_tit:
            novo_titulo = components.text_input(
                label=f"titulo_lbl_{ctx['id']}",
                value=ctx['titulo'],
                key=f"ren_title_input_{ctx['id']}",
                label_visibility="collapsed"
            )
        with c_btn_tit:
            if components.button_secondary("Aplicar Título", key=f"btn_aplicar_tit_{ctx['id']}"):
                ctx['titulo'] = novo_titulo
                st.rerun()
                
        with c_xml_imp:
            xml_file = st.file_uploader("Importar XML", type=["xml"], key=f"xml_import_{ctx['id']}", label_visibility="collapsed")
            if xml_file:
                # --- TRAVA ANTI-LOOP: Evita que o Streamlit reimporte o mesmo arquivo XML repetidamente ---
                if ctx.get('last_imported_xml_id') != xml_file.file_id:
                    # Mecanismo Visual de Barra de Progresso do XML
                    p_bar_xml = st.progress(0)
                    txt_status_xml = st.empty()
                    etapas_xml = [
                        ("Carregando arquivo XML...", 0.2),
                        ("Validando integridade dos nós...", 0.5),
                        ("Mapeando histórico e tabelas relacionais...", 0.8),
                        ("Processamento XML concluído!", 1.0)
                    ]
                    for msg_xml, prog_xml in etapas_xml:
                        txt_status_xml.markdown(f"**Status:** {msg_xml}")
                        p_bar_xml.progress(prog_xml)
                        time.sleep(0.15)
                    
                    ctx_atualizado = importar_contratacao_xml(xml_file, ctx['id'])
                    if ctx_atualizado:
                        ctx_atualizado['last_imported_xml_id'] = xml_file.file_id
                        st.session_state['lista_contratacoes'][idx_c] = ctx_atualizado
                        st.success("Dados da demanda importados via XML!")
                        time.sleep(1)
                        st.rerun()
                        
        with c_xml_exp:
            xml_data = exportar_contratacao_xml(ctx)
            st.download_button(
                label="Exportar XML",
                data=xml_data,
                file_name=f"Demanda_{ctx['titulo'].replace(' ', '_')}.xml",
                mime="application/xml",
                key=f"xml_export_{ctx['id']}",
                use_container_width=True
            )
        
        layout.divider_sub()
        
        # --- BLOCO EXPAND/COLLAPSE: OBJETO (SEM EMOJIS) ---
        with st.expander(f"{get_text('GRP_OBJETO')}", expanded=True):
            col_obj, col_itens = st.columns([1, 2], gap="large")
            if not ctx['tr_objeto_salvo']:
                with col_obj:
                    st.markdown(f"{ESTILO_LABEL_PADRAO}{get_text('LBL_COL_ESQ')}</label>", unsafe_allow_html=True)
                    txt_input = components.text_area(value=ctx['objeto_contratacao'], placeholder=get_text("LBL_DICA_COLA"), key=f"obj_area_{ctx['id']}")
                
                with col_itens:
                    st.markdown(f"{ESTILO_LABEL_PADRAO}{get_text('LBL_COL_DIR')}</label>", unsafe_allow_html=True)
                    df_tr_editado = components.data_editor(ctx['df_tr'], column_config=CONFIG_COLUNAS_TABELA, key=f"editor_tr_{ctx['id']}")
                
                with col_obj:
                    c_btn1, c_btn2 = st.columns([1, 1], gap="medium")
                    with c_btn1:
                        if components.button_primary(get_text("BTN_SALVAR_OBJ"), key=f"save_obj_btn_{ctx['id']}"):
                            if txt_input.strip():
                                ctx['objeto_contratacao'] = txt_input.strip()
                            df_validos = df_tr_editado.dropna(subset=["Item", "Descrição"]).copy()
                            df_validos['Lote'] = df_validos['Lote'].replace(r'^\s*$', pd.NA, regex=True).ffill()
                            df_validos['Quantidade'] = df_validos['Quantidade'].apply(parse_numero_localizado)
                            
                            ctx['df_tr'] = df_validos
                            ctx['tr_objeto_salvo'] = True
                            ctx['tr_itens_salvos'] = True
                            
                            for _, row in df_validos.iterrows():
                                h = gerar_hash_item(row)
                                if h not in ctx['banco_precos']:
                                    ctx['banco_precos'][h] = {"df_pncp": pd.DataFrame(columns=cols_pncp), "df_manual_rastreio": pd.DataFrame(columns=cols_rastreio), "historico_buscas": pd.DataFrame(columns=cols_historico_busca), "estatistica_pronta": False, "media_saneada": 0.0, "mediana": 0.0, "amostras": 0, "df_validos": pd.DataFrame(), "df_outliers": pd.DataFrame()}
                            st.rerun()
                                
                    with c_btn2:
                        up_file_tr = components.file_uploader(key=f"up_{ctx['id']}")
                        if up_file_tr:
                            # Mecanismo Visual de Barra de Progresso do Arquivo da Demanda
                            p_bar_tr = st.progress(0)
                            txt_status_tr = st.empty()
                            etapas_tr = [
                                ("Iniciando motor avançado Smart Extract...", 0.2),
                                ("Fazendo varredura semântica de marcadores...", 0.5),
                                ("Isolando descrição e minerando tabelas de itens...", 0.8),
                                ("Processamento estrutural finalizado!", 1.0)
                            ]
                            for msg_tr, prog_tr in etapas_tr:
                                txt_status_tr.markdown(f"**Status:** {msg_tr}")
                                p_bar_tr.progress(prog_tr)
                                time.sleep(0.15)
                                
                            obj_texto, df_extraido = processar_arquivo_inteligente(up_file_tr)
                            if obj_texto or not df_extraido.empty:
                                if obj_texto: 
                                    ctx['objeto_contratacao'] = obj_texto
                                ctx['tr_objeto_salvo'] = True
                                if not df_extraido.empty:
                                    ctx['df_tr'] = df_extraido
                                    ctx['tr_itens_salvos'] = True
                                    
                                for _, row in df_extraido.iterrows():
                                    h = gerar_hash_item(row)
                                    if h not in ctx['banco_precos']:
                                        ctx['banco_precos'][h] = {"df_pncp": pd.DataFrame(columns=cols_pncp), "df_manual_rastreio": pd.DataFrame(columns=cols_rastreio), "historico_buscas": pd.DataFrame(columns=cols_historico_busca), "estatistica_pronta": False, "media_saneada": 0.0, "mediana": 0.0, "amostras": 0, "df_validos": pd.DataFrame(), "df_outliers": pd.DataFrame()}
                                st.success(get_text("MSG_SUCESSO_IMPORT"))
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.error("Falha estrutural: Nenhuma descrição (1. OBJETO) ou tabela de itens válida foi reconhecida no arquivo.")
            else:
                with col_obj:
                    st.markdown(f"{ESTILO_LABEL_PADRAO}{get_text('LBL_COL_ESQ')}</label>", unsafe_allow_html=True)
                    st.info(ctx['objeto_contratacao'])
                    if components.button_secondary(get_text("BTN_EDITAR_OBJ"), key=f"edt_obj_btn_{ctx['id']}"):
                        ctx['tr_objeto_salvo'] = False
                        ctx['tr_itens_salvos'] = False
                        st.rerun()
                with col_itens:
                    st.markdown(f"{ESTILO_LABEL_PADRAO}{get_text('LBL_COL_DIR')}</label>", unsafe_allow_html=True)
                    components.dataframe(ctx['df_tr'].dropna(subset=["Item", "Descrição"]), column_config=CONFIG_COLUNAS_TABELA)

        layout.spacing()

        # --- BLOCO EXPAND/COLLAPSE: ANÁLISE DE MERCADO (SEM EMOJIS) ---
        with st.expander(f"{get_text('GRP_ANALISE')}", expanded=ctx['tr_itens_salvos']):
            if ctx['tr_itens_salvos']:
                df_validos_tr = ctx['df_tr'].dropna(subset=["Item", "Descrição"])
                qtd_itens = len(df_validos_tr)
                delay_sugerido = 1 if qtd_itens <= 3 else (3 if qtd_itens <= 10 else 5)
                if ctx['delay_pncp'] is None: ctx['delay_pncp'] = delay_sugerido
                
                layout.title_section(get_text('TITULO_PARAMETROS'))
                c_p1, c_p2, c_p3, c_p4 = st.columns(4, gap="small")
                with c_p1: ctx['regra_calculo'] = components.selectbox(get_text("LBL_REGRA_CALCULO"), ["Preços válidos - Mediana ±25% e Média"], key=f"regra_calc_{ctx['id']}")
                with c_p2: ctx['meses_corte'] = components.slider(get_text("LBL_MESES_CORTE"), min_value=12, max_value=60, value=ctx['meses_corte'], step=6, format_str="%d meses", key=f"corte_{ctx['id']}")
                with c_p3: ctx['paginas_pncp'] = components.number_input(get_text("LBL_PAGINAS_PNCP"), min_value=1, max_value=5, value=ctx['paginas_pncp'], key=f"pag_{ctx['id']}")
                with c_p4: ctx['delay_pncp'] = components.selectbox(get_text("LBL_DELAY_ANTI_BOT"), options=[1, 2, 3, 5, 10], index=[1, 2, 3, 5, 10].index(ctx['delay_pncp']), format_func=lambda x: f"{x}s" + (" (Risco WAF)" if x==1 else (" (Recomendado)" if x==delay_sugerido else "")), key=f"delay_{ctx['id']}")
                
                layout.divider_main()
                
                # --- PROCESSADOR DE EXECUTOR EM LOTE ---
                is_massa_active = ctx['massa_status'] != 'idle'
                if not is_massa_active:
                    if components.button_primary(get_text("BTN_BUSCAR_MASSA"), key=f"btn_massa_{ctx['id']}"):
                        ctx['massa_status'] = 'running'
                        ctx['massa_idx'] = 0
                        ctx['massa_novos'] = 0
                        ctx['massa_total'] = len(df_validos_tr)
                        st.rerun()
                else:
                    col_ctrl1, col_ctrl2 = st.columns(2)
                    with col_ctrl1:
                        if ctx['massa_status'] == 'running':
                            if components.button_secondary(get_text("BTN_PAUSAR_MASSA"), key=f"btn_pause_{ctx['id']}"): ctx['massa_status'] = 'paused'; st.rerun()
                        elif ctx['massa_status'] == 'paused':
                            if components.button_primary(get_text("BTN_RETOMAR_MASSA"), key=f"btn_resume_{ctx['id']}"): ctx['massa_status'] = 'running'; st.rerun()
                    with col_ctrl2:
                        if components.button_secondary(get_text("BTN_PARAR_MASSA"), key=f"btn_stop_{ctx['id']}"): ctx['massa_status'] = 'idle'; ctx['massa_idx'] = 0; st.rerun()
                    
                    st.progress(ctx['massa_idx'] / ctx['massa_total'] if ctx['massa_total'] > 0 else 0)
                    if ctx['massa_status'] == 'paused': st.warning(get_text("STATUS_MASSA_PAUSADO").format(atual=ctx['massa_idx']+1, total=ctx['massa_total']))
                    
                    if ctx['massa_status'] == 'running':
                        if ctx['massa_idx'] < ctx['massa_total']:
                            idx = ctx['massa_idx']
                            r_m = df_validos_tr.iloc[idx]
                            h_m = gerar_hash_item(r_m)
                            
                            if h_m not in ctx['palavras_chave_massa'] or not ctx['palavras_chave_massa'][h_m].strip():
                                ctx['palavras_chave_massa'][h_m] = " ".join(r_m['Descrição'].split()[:3])
                        
                            termo_m = ctx['palavras_chave_massa'][h_m]
                            banco_m = ctx['banco_precos'][h_m]
                            st.info(get_text("STATUS_MASSA_ITEM_BUSCA").format(atual=idx+1, total=ctx['massa_total'], termo=termo_m))
                            
                            engine_m = PNCPEngine()
                            editais_m, tipo_m = engine_m.buscar_editais_inteligente(termo_m, ctx['paginas_pncp'], None)
                            all_items_m = []
                            if editais_m:
                                ext_status = st.empty()
                                with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor_m:
                                    futures_m = [executor_m.submit(engine_m.minerar_itens, ed_m, termo_m) for ed_m in editais_m]
                                for ed_idx, f_m in enumerate(concurrent.futures.as_completed(futures_m), 1):
                                    all_items_m.extend(f_m.result())
                                ext_status.empty()
                            
                            ctx['massa_novos'] += len(all_items_m)
                            banco_m["historico_buscas"] = pd.concat([banco_m["historico_buscas"], pd.DataFrame([{"Data/Hora": datetime.now(fuso_br).strftime("%d/%m/%Y %H:%M"), "Termo Pesquisado": termo_m, "Novos Registros": len(all_items_m)}])], ignore_index=True)
                            if all_items_m:
                                df_novos_m = pd.DataFrame(all_items_m)
                                df_novos_m.insert(0, "Válido?", True)
                                banco_m["df_pncp"] = pd.concat([banco_m["df_pncp"], df_novos_m], ignore_index=True)
                            
                            ctx['massa_idx'] += 1
                            if ctx['massa_idx'] < ctx['massa_total']:
                                delay_segundos = int(ctx['delay_pncp'])
                                st.warning(get_text("STATUS_MASSA_PAUSA").format(atual=idx+1, total=ctx['massa_total'], seg=delay_segundos))
                                time.sleep(delay_segundos); st.rerun()
                            else:
                                ctx['massa_status'] = 'idle'
                                st.success(get_text("STATUS_MASSA_FIM").format(qtd=ctx['massa_novos']))
                                time.sleep(3); st.rerun()
                
                layout.spacing()
                components.render_table_header(get_text("COL_LOTE"), get_text("COL_DESC"), get_text("COL_STATUS"), get_text("COL_ACOES"))

                for _, row in df_validos_tr.iterrows():
                    h_id = gerar_hash_item(row)
                    banco = ctx['banco_precos'].get(h_id)
                    if not banco: continue
                    
                    if h_id not in ctx['palavras_chave_massa']:
                        ctx['palavras_chave_massa'][h_id] = " ".join(row['Descrição'].split()[:3])
                        
                    lote_lbl = row['Lote'] if pd.notna(row['Lote']) and str(row['Lote']).strip() != "" else "Único"
                    st.markdown(components.render_item_row_start(), unsafe_allow_html=True)
                    c1, c2, c3, c4 = st.columns([1, 4, 1.5, 3], gap="small")
                    c1.write(f"**{row['Item']}** ({lote_lbl})")
                    
                    with c2:
                        st.markdown(f"<span>{row['Descrição']}</span>", unsafe_allow_html=True)
                        ctx['palavras_chave_massa'][h_id] = components.text_input(
                            label=f"kw_lbl_{ctx['id']}_{h_id}",
                            value=ctx['palavras_chave_massa'][h_id],
                            key=f"kw_{ctx['id']}_{h_id}",
                            label_visibility="collapsed"
                        )
                        
                    c3.markdown(components.render_status_badge(banco['estatistica_pronta'], formatar_moeda_simples(banco['media_saneada']), get_text("BADGE_PENDENTE")), unsafe_allow_html=True)
                    with c4:
                        b1, b2, b3 = st.columns(3, gap="small")
                        with b1:
                            if components.button_secondary(get_text("BTN_PNCP"), key=f"p_{ctx['id']}_{h_id}"): ctx['acao_ativa'] = ("pncp", h_id); st.rerun()
                        with b2:
                            if components.button_secondary(get_text("BTN_MANUAL"), key=f"c_{ctx['id']}_{h_id}"): ctx['acao_ativa'] = ("manual", h_id); st.rerun()
                        with b3:
                            if components.button_secondary(get_text("BTN_VALIDAR"), key=f"v_{ctx['id']}_{h_id}"): ctx['acao_ativa'] = ("validar", h_id); st.rerun()
                    st.markdown("</div>", unsafe_allow_html=True)

                    # --- SUB-FORMULÁRIOS INLINE POR CONTEXTO DE ITEM ---
                    acao, active_hash = ctx['acao_ativa']
                    if acao and active_hash == h_id:
                        st.markdown("<div style='padding: 15px; background-color: #FFFFFF; border-left: 4px solid #0F2C4C; border-radius: 4px; margin-bottom: 20px; box-shadow: 0 1px 3px rgba(0,0,0,0.05);'>", unsafe_allow_html=True)
                        
                        if acao == "pncp":
                            with st.form(f"form_pncp_{ctx['id']}_{h_id}"):
                                termo = components.text_input("Busca:", value=ctx['palavras_chave_massa'][h_id], key=f"ti_{ctx['id']}_{h_id}")
                                if components.form_submit_button(get_text("BTN_BUSCAR_PNCP")):
                                    engine = PNCPEngine()
                                    status = st.status(get_text("STATUS_BUSCA_INICIAL"), expanded=True)
                                    editais, tipo = engine.buscar_editais_inteligente(termo, ctx['paginas_pncp'], status)
                                    all_items = []
                                    if editais:
                                        total_ed = len(editais)
                                        status.update(label=get_text("STATUS_EXTRAINDO").format(tipo=tipo), state="running")
                                        progress_bar = st.progress(0)
                                        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                                            futures = [executor.submit(engine.minerar_itens, ed, termo) for ed in editais]
                                        for i, f in enumerate(concurrent.futures.as_completed(futures), 1):
                                            all_items.extend(f.result())
                                        progress_bar.progress(i / total_ed)
                                        progress_bar.empty()
                                        
                                    qtd = len(all_items)
                                    banco["historico_buscas"] = pd.concat([banco["historico_buscas"], pd.DataFrame([{"Data/Hora": datetime.now(fuso_br).strftime("%d/%m/%Y %H:%M"), "Termo Pesquisado": termo, "Novos Registros": qtd}])], ignore_index=True)
                                    if all_items:
                                        df_novos = pd.DataFrame(all_items)
                                        df_novos.insert(0, "Válido?", True)
                                        banco["df_pncp"] = pd.concat([banco["df_pncp"], df_novos], ignore_index=True)
                                        status.update(label=get_text("STATUS_CONCLUIDO").format(qtd=qtd), state="complete")
                                    else:
                                        status.update(label=get_text("MSG_ERRO_NADA"), state="error")
                            components.data_editor(banco["historico_buscas"], column_config={"Termo Pesquisado": st.column_config.TextColumn("Termo Pesquisado", width="large")}, key=f"hist_{ctx['id']}_{h_id}")

                        elif acao == "manual":
                            with st.form(f"form_man_{ctx['id']}_{h_id}"):
                                layout.title_section(get_text("LBL_NOVO_REGISTRO"))
                                c1_m, c2_m, c3_m = st.columns(3)
                                with c1_m: m_emp = components.text_input("Empresa", value="", key=f"emp_{ctx['id']}_{h_id}")
                                with c2_m: m_preco = components.number_input("Preço", min_value=0.0, key=f"preco_{ctx['id']}_{h_id}")
                                with c3_m: m_sit = components.selectbox("Situação:", opcoes_situacao, key=f"sit_{ctx['id']}_{h_id}")
                                if components.form_submit_button(get_text("BTN_SALVAR_REGISTRO")):
                                    log = {"Data do Contato": datetime.now(fuso_br).strftime("%d/%m/%Y"), "Horário": datetime.now(fuso_br).strftime("%H:%M"), "Empresa": m_emp, "Situação": m_sit, "Preço": m_preco, "Valor Unitário": formatar_moeda_ordenavel(m_preco), "Link da fonte": "", "Descrição da fonte": ""}
                                    banco["df_manual_rastreio"] = pd.concat([banco["df_manual_rastreio"], pd.DataFrame([log])], ignore_index=True)
                                    st.rerun()
                            components.data_editor(banco["df_manual_rastreio"], column_config={"Link da fonte": st.column_config.LinkColumn("Link")}, key=f"man_{ctx['id']}_{h_id}")

                        elif acao == "validar":
                            df_pncp_atual = banco["df_pncp"].copy()
                            if not df_pncp_atual.empty:
                                limite_data = (datetime.now() - relativedelta(months=ctx['meses_corte']))
                                df_pncp_atual['Válido?'] = pd.to_datetime(df_pncp_atual['Data'], format="%d/%m/%Y", errors='coerce').apply(lambda d: True if pd.isna(d) else d >= limite_data)
                            
                            df_man_valido = pd.DataFrame()
                            if not banco["df_manual_rastreio"].empty:
                                df_man_valido = banco["df_manual_rastreio"].copy()
                                df_man_valido["Origem"] = "Manual"; df_man_valido["Item"] = row['Descrição'];
                                df_man_valido["Qtd"] = 1; df_man_valido["Tipo"] = "Manual"; df_man_valido["Válido?"] = True
                                df_man_valido = df_man_valido.rename(columns={"Data do Contato": "Data", "Empresa": "Empresa/Órgão"})[[c for c in cols_pncp if c in df_man_valido.columns]]
                            
                            with st.form(f"form_val_{ctx['id']}_{h_id}"):
                                col_config_val = {"Empresa/Órgão": st.column_config.TextColumn("Empresa/Órgão", disabled=True), "Origem": st.column_config.LinkColumn("Origem", disabled=True), "Item": st.column_config.TextColumn("Item", disabled=True), "Preço": st.column_config.NumberColumn("Preço", format="R$ %.2f", disabled=True)}
                                pncp_res = components.data_editor(df_pncp_atual, column_config=col_config_val, key=f"val_pncp_ed_{ctx['id']}_{h_id}") if not df_pncp_atual.empty else pd.DataFrame()
                                man_res = components.data_editor(df_man_valido, column_config=col_config_val, key=f"val_man_ed_{ctx['id']}_{h_id}") if not df_man_valido.empty else pd.DataFrame()
                                
                                if components.form_submit_button(get_text("BTN_CALCULAR_MEDIANA")):
                                    df_merge = pd.concat([pncp_res, man_res], ignore_index=True)
                                    if not df_merge.empty and "Válido?" in df_merge.columns:
                                        df_base_calculo = df_merge[df_merge["Válido?"] == True]
                                    else:
                                        df_base_calculo = pd.DataFrame(columns=["Preço"])
                                        
                                    df_v, df_o, m_geral, _, _ = processar_precos_regra(df_base_calculo, ctx['regra_calculo'])
                                    banco["df_validos"] = ordenar_validos(df_v); banco["df_outliers"] = ordenar_outliers(df_o); banco["media_saneada"] = df_v['Preço'].mean() if not df_v.empty else 0; banco["mediana"] = m_geral; banco["amostras"] = len(df_v); banco["estatistica_pronta"] = True
                                    ctx['acao_ativa'] = (None, None); st.rerun()

                            if banco["estatistica_pronta"]:
                                layout.spacing()
                                maior_v = banco["df_validos"]['Preço'].max() if not banco["df_validos"].empty else 0
                                l1_c1, l1_c2, l1_c3 = st.columns(3)
                                l1_c1.markdown(components.render_metric_card(get_text('LBL_PRECO_FINAL'), formatar_moeda_simples(banco['media_saneada'])), unsafe_allow_html=True)
                                l1_c2.markdown(components.render_metric_card(get_text('LBL_AMOSTRAS'), banco['amostras']), unsafe_allow_html=True)
                                l1_c3.markdown(components.render_metric_card(get_text('LBL_MAIOR_VALOR'), formatar_moeda_simples(maior_v)), unsafe_allow_html=True)
                        
                        st.markdown("</div>", unsafe_allow_html=True)

        layout.spacing()

        # --- BLOCO EXPAND/COLLAPSE: RELATÓRIO (SEM EMOJIS) ---
        with st.expander(f"{get_text('GRP_RELATORIO')}", expanded=ctx['tr_itens_salvos']):
            if ctx['tr_itens_salvos']:
                df_validos_tr = ctx['df_tr'].dropna(subset=["Item", "Descrição"])
                
                lotes_dict = {}
                valor_total_global = 0.0
                for _, row in df_validos_tr.iterrows():
                    l_key = row["Lote"] if pd.notna(row["Lote"]) and str(row["Lote"]).strip() else "Único"
                    if l_key not in lotes_dict: lotes_dict[l_key] = []
                    banco = ctx['banco_precos'].get(gerar_hash_item(row))
                    if banco:
                        med = banco['media_saneada']
                        q = float(row.get("Quantidade", 1))
                        valor_total_global += (med * q)
                        lotes_dict[l_key].append({"Item": row["Item"], "Descrição": row["Descrição"], "Qtd": row["Quantidade"], "Unid.": row["Métrica"], "Preço Numérico": med})
                
                components.render_total_global(get_text('VALOR_TOTAL_GLOBAL'), formatar_moeda_simples(valor_total_global))

                if components.button_primary(get_text("BTN_GERAR_PDF"), key=f"pdf_g_{ctx['id']}"):
                    pdf_bytes = gerar_pdf_oficial(df_validos_tr, lotes_dict, valor_total_global, ctx['objeto_contratacao'], ctx['regra_calculo'], datetime.now(fuso_br).strftime('%d/%m/%Y %H:%M'), ctx['banco_precos'])
                    st.download_button(get_text("BTN_SALVAR_PDF"), data=pdf_bytes, file_name=f"Analise_Mercado_{ctx['titulo'].replace(' ', '_')}.pdf", mime="application/pdf", key=f"pdf_d_{ctx['id']}")

        layout.spacing()

layout.divider_main()
if components.button_primary(get_text("BTN_NOVA_CONTRATACAO"), key="btn_new_contract"):
    proximo_id = len(st.session_state['lista_contratacoes']) + 1
    st.session_state['lista_contratacoes'].append(
        {
            "id": proximo_id,
            "titulo": f"Contratação {proximo_id}",
            "tr_objeto_salvo": False,
            "tr_itens_salvos": False,
            "objeto_contratacao": "",
            "df_tr": pd.DataFrame(columns=["Lote", "Item", "Descrição", "Métrica", "Tipo", "Quantidade"]),
            "banco_precos": {},
            "acao_ativa": (None, None),
            "regra_calculo": "Preços válidos - Mediana ±25% e Média",
            "meses_corte": 24,
            "paginas_pncp": 3,
            "delay_pncp": None,
            "palavras_chave_massa": {},
            "massa_status": "idle",
            "massa_idx": 0,
            "massa_total": 0,
            "massa_novos": 0
        }
    )
    st.rerun()